import io
import streamlit as st
from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

# ==========================================
# 1. CẤU HÌNH TRANG WEB STREAMLIT
# ==========================================
st.set_page_config(page_title="Tạo Văn Bản Chuẩn NĐ 30", page_icon="📄", layout="centered")

def hide_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(r'<w:tblBorders %s><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/><w:insideH w:val="none"/><w:insideV w:val="none"/></w:tblBorders>' % nsdecls('w'))
    tblPr.append(borders)

def format_text_run(run, font_name="Times New Roman", size_pt=13, bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.extend([fldChar1, instrText, fldChar2, fldChar3])

# ==========================================
# 2. HÀM TẠO FILE DOCX ĐÚNG CHUẨN NĐ 30
# ==========================================
def generate_nd30_docx(chu_quan, ban_hanh, text_content):
    doc = Document()
    
    # Định dạng Khổ giấy A4 và Lề trang
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Cm(2.0)    
        section.bottom_margin = Cm(2.0) 
        section.left_margin = Cm(3.0)   
        section.right_margin = Cm(2.0)  

        section.different_first_page_header_footer = True 
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        run_page = header_para.add_run()
        add_page_number(run_page)
        format_text_run(run_page, size_pt=13, bold=False)

    # --- PHẦN ĐẦU TRANG ---
    table = doc.add_table(rows=1, cols=2)
    hide_table_borders(table)
    table.columns[0].width = Cm(7.0)
    table.columns[1].width = Cm(9.0)
    
    cell_left = table.cell(0, 0)
    cell_right = table.cell(0, 1)
    
    p_left_1 = cell_left.paragraphs[0]
    p_left_1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_left_1.paragraph_format.line_spacing = 1.0
    p_left_1.paragraph_format.space_after = Pt(0)
    
    if chu_quan.strip():
        run_cq = p_left_1.add_run(chu_quan.strip().upper() + "\n")
        format_text_run(run_cq, size_pt=12, bold=False)
        
    run_bh = p_left_1.add_run(ban_hanh.strip().upper())
    format_text_run(run_bh, size_pt=13, bold=True)
    
    p_left_line = cell_left.add_paragraph()
    p_left_line.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_left_line.paragraph_format.line_spacing = 1.0
    p_left_line.paragraph_format.space_after = Pt(2)
    line_len_left = max(6, int(len(ban_hanh.strip()) * 0.35))
    run_line_left = p_left_line.add_run("_" * line_len_left)
    format_text_run(run_line_left, size_pt=11, bold=True)

    p_left_so = cell_left.add_paragraph()
    p_left_so.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_left_so.paragraph_format.line_spacing = 1.0
    run_so = p_left_so.add_run("Số: ....../...............")
    format_text_run(run_so, size_pt=13, bold=False)

    p_right_qh = cell_right.paragraphs[0]
    p_right_qh.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_right_qh.paragraph_format.line_spacing = 1.0
    p_right_qh.paragraph_format.space_after = Pt(0)
    
    run_qh = p_right_qh.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    format_text_run(run_qh, size_pt=13, bold=True)
    
    p_right_tn = cell_right.add_paragraph()
    p_right_tn.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_right_tn.paragraph_format.line_spacing = 1.0
    p_right_tn.paragraph_format.space_after = Pt(0)
    run_tn = p_right_tn.add_run("Độc lập - Tự do - Hạnh phúc")
    format_text_run(run_tn, size_pt=14, bold=True)
    
    p_right_line = cell_right.add_paragraph()
    p_right_line.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_right_line.paragraph_format.line_spacing = 1.0
    p_right_line.paragraph_format.space_after = Pt(2)
    run_line_right = p_right_line.add_run("________________________")
    format_text_run(run_line_right, size_pt=11, bold=True)

    p_right_date = cell_right.add_paragraph()
    p_right_date.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_right_date.paragraph_format.line_spacing = 1.0
    run_date = p_right_date.add_run("Rạch Giá, ngày ... tháng ... năm ...")
    format_text_run(run_date, size_pt=13, bold=False, italic=True)

    # ==========================================
    # PHẦN THÂN VĂN BẢN (Tôn trọng tuyệt đối các lần xuống dòng)
    # ==========================================
    p_space1 = doc.add_paragraph()
    p_space1.paragraph_format.space_before = Pt(12)

    # Tách dòng đúng theo số lần bạn ấn Enter
    paragraphs = text_content.split('\n')
    for p_text in paragraphs:
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY 
        p.paragraph_format.line_spacing = 1.15       
        p.paragraph_format.space_after = Pt(6)       

        # Nếu là dòng trống (do ấn Enter cách đoạn), giữ nguyên khoảng trống đó
        if p_text.strip() == "":
            continue
            
        # Nếu có chữ, thụt đầu dòng 1cm và định dạng cỡ 13
        p.paragraph_format.first_line_indent = Cm(1.0) 
        run_body = p.add_run(p_text.strip())
        format_text_run(run_body, size_pt=13, bold=False)

    # ==========================================
    # PHẦN CUỐI TRANG: NƠI NHẬN & CHỮ KÝ
    # ==========================================
    p_space2 = doc.add_paragraph()
    p_space2.paragraph_format.space_before = Pt(6)

    table_footer = doc.add_table(rows=1, cols=2)
    hide_table_borders(table_footer)
    table_footer.columns[0].width = Cm(7.0)
    table_footer.columns[1].width = Cm(9.0)
    
    cell_f_left = table_footer.cell(0, 0)
    cell_f_right = table_footer.cell(0, 1)
    
    p_nn_title = cell_f_left.paragraphs[0]
    p_nn_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p_nn_title.paragraph_format.line_spacing = 1.0
    p_nn_title.paragraph_format.space_after = Pt(0)
    run_nn_title = p_nn_title.add_run("Nơi nhận:")
    format_text_run(run_nn_title, size_pt=12, bold=True, italic=True)
    
    p_nn_1 = cell_f_left.add_paragraph()
    p_nn_1.paragraph_format.line_spacing = 1.0
    p_nn_1.paragraph_format.space_after = Pt(0)
    run_nn_1 = p_nn_1.add_run("- ......;")
    format_text_run(run_nn_1, size_pt=11, bold=False)
    
    p_nn_2 = cell_f_left.add_paragraph()
    p_nn_2.paragraph_format.line_spacing = 1.0
    p_nn_2.paragraph_format.space_after = Pt(0)
    run_nn_2 = p_nn_2.add_run("- Lưu: ......")
    format_text_run(run_nn_2, size_pt=11, bold=False)

    p_cv = cell_f_right.paragraphs[0]
    p_cv.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_cv.paragraph_format.line_spacing = 1.0
    p_cv.paragraph_format.space_after = Pt(0)
    run_cv = p_cv.add_run("TRƯỞNG PHÒNG")
    format_text_run(run_cv, size_pt=13, bold=True)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==========================================
# 3. GIAO DIỆN NGƯỜI DÙNG STREAMLIT
# ==========================================
st.title("📄 Ứng Dụng Chuẩn Hóa Văn Bản Hành Chính")
st.markdown("Chuyển đổi nội dung văn bản (từ AI) thành file Word (`.docx`) đúng định dạng **Nghị định 30/2020/NĐ-CP**.")

st.subheader("1. Thông tin cơ quan")
col1, col2 = st.columns(2)
with col1:
    input_chu_quan = st.text_input(
        "Cơ quan chủ quản (nếu có):", 
        placeholder="Ví dụ: UBND TỈNH KIÊN GIANG"
    )
with col2:
    input_ban_hanh = st.text_input(
        "Cơ quan ban hành văn bản *", 
        placeholder="Ví dụ: SỞ THÔNG TIN VÀ TRUYỀN THÔNG"
    )

st.info("💡 Các thông tin cố định như Quốc hiệu, Tiêu ngữ, Số hiệu, Đánh số trang (không đánh trang 1), Nơi nhận, Chức danh người ký sẽ được tạo tự động.")

st.subheader("2. Nội dung văn bản")
user_input = st.text_area("Dán nội dung phần thân văn bản vào đây:", height=300)

if st.button("⚡ Tiến hành tạo file Word"):
    if not input_ban_hanh.strip():
        st.error("Vui lòng điền 'Cơ quan ban hành văn bản' trước khi tiếp tục.")
    elif not user_input.strip():
        st.warning("Vui lòng nhập nội dung thân văn bản.")
    else:
        with st.spinner("Đang xử lý định dạng..."):
            docx_file = generate_nd30_docx(input_chu_quan, input_ban_hanh, user_input)
            
            st.success("🎉 Đã chuẩn hóa định dạng thành công!")
            st.download_button(
                label="⬇️ Tải xuống file Văn_bản_ND30.docx",
                data=docx_file,
                file_name="Van_ban_ND30.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
